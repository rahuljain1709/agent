import io
from docx import Document
from openai import OpenAI

class Researcher:
    def __init__(self, api_key):
        self.client = OpenAI(api_key=api_key)

    def summarize_article(self, title, url, content):
        prompt = f"""
        Summarize the article in bullet points and one paragraph.
        Title: {title}
        URL: {url}
        Content: {content}
        """

        resp = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":prompt}]
        )

        return resp.choices[0].message.content

    def build_document(self, query, results):
        doc = Document()
        doc.add_heading(f"Research report: {query}", 1)

        for r in results:
            summary = self.summarize_article(
                r["title"], r["url"], r["content"]
            )
            doc.add_heading(r["title"], 2)
            doc.add_paragraph(f"URL: {r['url']}")
            doc.add_paragraph(summary)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        filename = f"{query.replace(' ','_')}.docx"
        return bio.read(), filename
